"""
Word document builder.
Imports charts from app.report.charts — no direct matplotlib here.
"""
from __future__ import annotations
import io
import logging
from datetime import datetime
from typing import Optional

import app.report.charts as C    # direct submodule — avoids circular __init__ risk

logger = logging.getLogger(__name__)

try:
    from docx import Document
    from docx.enum.text  import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns    import qn
    from docx.oxml       import OxmlElement
    from docx.shared     import Inches, Pt, RGBColor, Cm
    _OK = True
except ImportError:
    _OK = False
    logger.warning("python-docx not installed.")


# ── Helpers ───────────────────────────────────────────────────────────────────

def _cell_bg(cell, hex_col: str) -> None:
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_col.lstrip("#").upper())
    tcPr.append(shd)


def _header_row(table, bg="1E4D8C") -> None:
    for cell in table.rows[0].cells:
        _cell_bg(cell, bg)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.size = Pt(9)


def _kv_table(doc, rows, widths=(2.5, 3.5)) -> None:
    tbl = doc.add_table(rows=len(rows)+1, cols=2)
    tbl.style = "Table Grid"
    tbl.rows[0].cells[0].text = "Metric"
    tbl.rows[0].cells[1].text = "Value"
    _header_row(tbl)
    for i, (k, v) in enumerate(rows, 1):
        tbl.rows[i].cells[0].text = k
        tbl.rows[i].cells[1].text = str(v)
        if i % 2 == 0:
            for c in tbl.rows[i].cells: _cell_bg(c, "EBF0FA")
    for row in tbl.rows:
        row.cells[0].width = Inches(widths[0])
        row.cells[1].width = Inches(widths[1])


def _fmt(v, suf="", dec=1) -> str:
    if v is None: return "N/A"
    try: return f"{float(v):,.{dec}f}{suf}"
    except: return str(v)


def _bullets(doc, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def _chart(doc, buf: Optional[io.BytesIO], w=6.0) -> None:
    if buf:
        doc.add_picture(buf, width=Inches(w))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def _rec_color(rec: str) -> RGBColor:
    return {
        "APPROVE":     RGBColor(0x27, 0xAE, 0x60),
        "CONDITIONAL": RGBColor(0xE6, 0x7E, 0x22),
        "DECLINE":     RGBColor(0xC0, 0x39, 0x2B),
    }.get(rec, RGBColor(0x7F, 0x8C, 0x8D))


# ── Main builder ──────────────────────────────────────────────────────────────

def generate_word_report(analysis: dict, docs: list[dict]) -> bytes:
    if not _OK:
        raise RuntimeError("python-docx not installed. Run: pip install python-docx")

    doc     = Document()
    f       = analysis.get("financials", {})
    dec     = analysis.get("decision",   {})
    bd      = analysis.get("score_breakdown", {})
    az      = analysis.get("altman_z") or {}
    score   = analysis.get("credit_score", 0)
    rating  = analysis.get("credit_rating", "N/A")
    rec     = dec.get("recommendation", "N/A")
    rc      = _rec_color(rec)
    today   = datetime.now().strftime("%d %B %Y")
    company = f.get("company_name", "Unknown Company")
    curr    = f.get("currency", "")
    unit    = f.get("unit", "")

    for sec in doc.sections:
        sec.top_margin    = Cm(2.0); sec.bottom_margin = Cm(2.0)
        sec.left_margin   = Cm(2.5); sec.right_margin  = Cm(2.5)

    # ── Cover ─────────────────────────────────────────────────────────────────
    doc.add_paragraph(); doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("CREDIT ANALYSIS REPORT"); r.bold = True
    r.font.size = Pt(24); r.font.color.rgb = RGBColor(0x1E, 0x4D, 0x8C)
    doc.add_paragraph()
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(company); r2.bold = True; r2.font.size = Pt(18)
    doc.add_paragraph()
    for label, value in [("Fiscal Year", f.get("fiscal_year","N/A")),
                          ("Currency",   f"{curr} ({unit})"),
                          ("Report Date", today),
                          ("Analyst",    "Credit Analysis AI System")]:
        p3 = doc.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.add_run(f"{label}: ").bold = True; p3.add_run(value)
    doc.add_paragraph()
    pb = doc.add_paragraph(); pb.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rb = pb.add_run(f"CREDIT DECISION: {rec}  |  Score: {score}/100  ({rating})")
    rb.bold = True; rb.font.size = Pt(14); rb.font.color.rgb = rc
    doc.add_page_break()

    # ── 1. Executive Summary ─────────────────────────────────────────────────
    doc.add_heading("1. Executive Summary", 1)
    p = doc.add_paragraph()
    r = p.add_run(f"Decision: {rec}  |  Score: {score}/100 ({rating})  |  Altman Z′: {az.get('z_score','N/A')} ({az.get('zone','N/A')})")
    r.bold = True; r.font.size = Pt(12); r.font.color.rgb = rc
    doc.add_paragraph()
    if dec.get("narrative"): doc.add_paragraph(dec["narrative"])
    la = analysis.get("loan_amount_requested"); ml = dec.get("max_safe_loan")
    if la:
        p4 = doc.add_paragraph()
        p4.add_run("Loan Requested: ").bold = True
        p4.add_run(f"{la:,.0f} {curr}")
        if ml: p4.add_run(f"  |  Max Recommended: {ml:,.0f} {curr}")

    # ── 2. Company Overview ───────────────────────────────────────────────────
    doc.add_heading("2. Company Overview", 1)
    _kv_table(doc, [
        ("Company Name",     company),
        ("Ticker",           f.get("ticker") or "N/A"),
        ("Fiscal Year",      f.get("fiscal_year") or "N/A"),
        ("Reporting Period", f.get("reporting_period") or "N/A"),
        ("Currency / Unit",  f"{curr} / {unit}"),
        ("Sector",           f.get("sector") or "N/A"),
        ("Industry",         f.get("industry") or "N/A"),
    ])
    if f.get("qualitative_notes"):
        doc.add_paragraph()
        p5 = doc.add_paragraph(); p5.add_run("Business Context: ").bold = True
        p5.add_run(f.get("qualitative_notes"))

    # ── 3. Financial Analysis ─────────────────────────────────────────────────
    doc.add_heading("3. Financial Analysis", 1)
    _chart(doc, C.revenue_trend(f))
    doc.add_paragraph()
    doc.add_heading("3.1 Income Statement", 2)
    _kv_table(doc, [(k, _fmt(f.get(v), f" {curr}")) for k, v in [
        ("Revenue",              "revenue"),
        ("Revenue (Prior Year)", "revenue_prev"),
        ("Gross Profit",         "gross_profit"),
        ("EBITDA",               "ebitda"),
        ("Operating Profit",     "operating_profit"),
        ("Net Profit",           "net_profit"),
        ("Interest Expense",     "interest_expense"),
        ("D&A",                  "depreciation_amortization"),
    ] if f.get(v) is not None])
    doc.add_heading("3.2 Balance Sheet", 2)
    _kv_table(doc, [(k, _fmt(f.get(v), f" {curr}")) for k, v in [
        ("Current Assets",       "current_assets"),
        ("Total Assets",         "total_assets"),
        ("Current Liabilities",  "current_liabilities"),
        ("Total Liabilities",    "total_liabilities"),
        ("Net Assets / Equity",  "net_assets"),
        ("Cash",                 "cash"),
        ("Retained Earnings",    "retained_earnings"),
    ] if f.get(v) is not None])
    doc.add_heading("3.3 Cash Flow", 2)
    _kv_table(doc, [(k, _fmt(f.get(v), f" {curr}")) for k, v in [
        ("Operating Cash Flow", "operating_cash_flow"),
        ("CapEx",               "capex"),
        ("Free Cash Flow",      "free_cash_flow"),
    ] if f.get(v) is not None])
    segs = f.get("segments", [])
    if segs:
        doc.add_heading("3.4 Segments", 2)
        st = doc.add_table(rows=len(segs)+1, cols=3); st.style = "Table Grid"
        for j, h in enumerate(["Segment", f"Revenue ({curr})", f"Profit ({curr})"]):
            st.rows[0].cells[j].text = h
        _header_row(st)
        for i, seg in enumerate(segs, 1):
            st.rows[i].cells[0].text = seg.get("name","")
            st.rows[i].cells[1].text = _fmt(seg.get("revenue"))
            st.rows[i].cells[2].text = _fmt(seg.get("profit"))

    # ── 4. Credit Score ───────────────────────────────────────────────────────
    doc.add_page_break()
    doc.add_heading("4. Credit Score Analysis", 1)
    _chart(doc, C.score_breakdown(bd, score, rating))
    doc.add_paragraph()
    maxes = {"profitability":30,"leverage":25,"liquidity":20,"growth":15,"qualitative":10}
    st2 = doc.add_table(rows=len(bd)+2, cols=3); st2.style = "Table Grid"
    st2.rows[0].cells[0].text = "Category"
    st2.rows[0].cells[1].text = "Score"; st2.rows[0].cells[2].text = "Max"
    _header_row(st2)
    for i, cat in enumerate(["profitability","leverage","liquidity","growth","qualitative"], 1):
        st2.rows[i].cells[0].text = cat.title()
        st2.rows[i].cells[1].text = f"{bd.get(cat,0):.1f}"
        st2.rows[i].cells[2].text = str(maxes.get(cat,"?"))
        if i % 2 == 0:
            for c in st2.rows[i].cells: _cell_bg(c, "EBF0FA")
    tr = st2.rows[-1]
    tr.cells[0].text = "TOTAL"; tr.cells[1].text = f"{score:.1f}"; tr.cells[2].text = "100"
    for c in tr.cells: _cell_bg(c, "D6E4FF")
    for c in tr.cells:
        for p in c.paragraphs:
            for r in p.runs: r.bold = True
    if az:
        doc.add_paragraph()
        doc.add_heading("4.1 Altman Z′-Score", 2)
        _kv_table(doc, [
            ("Z′-Score", str(az.get("z_score","N/A"))),
            ("Zone",     az.get("zone","N/A")),
            ("X1 (Working Capital / Total Assets)", _fmt(az.get("x1"), dec=4)),
            ("X2 (Retained Earnings / Total Assets)", _fmt(az.get("x2"), dec=4)),
            ("X3 (EBIT / Total Assets)", _fmt(az.get("x3"), dec=4)),
            ("X4 (Book Equity / Total Liabilities)", _fmt(az.get("x4"), dec=4)),
            ("X5 (Revenue / Total Assets)", _fmt(az.get("x5"), dec=4)),
        ])
        pi = doc.add_paragraph()
        pi.add_run("Interpretation: ").bold = True
        pi.add_run("Z′ > 2.9 = Safe | 1.23–2.9 = Grey | < 1.23 = Distress").italic = True

    # ── 5. Ratio Analysis ─────────────────────────────────────────────────────
    doc.add_page_break()
    doc.add_heading("5. Financial Ratio Analysis", 1)
    _chart(doc, C.ratio_dashboard(f))
    doc.add_paragraph()
    ratio_rows = [
        ("Liquidity", None, None, None),
        ("Current Ratio",     _fmt(f.get("current_ratio"),"×"), "≥2.0×",   "Strong" if (f.get("current_ratio") or 0)>=2 else "Weak"),
        ("Quick Ratio",       _fmt(f.get("quick_ratio"),"×"),   "≥1.0×",   "Strong" if (f.get("quick_ratio") or 0)>=1 else "Weak"),
        ("Leverage", None, None, None),
        ("Equity Ratio",      _fmt(f.get("equity_ratio"),"%"),  "≥40%",    "Strong" if (f.get("equity_ratio") or 0)>=60 else ("OK" if (f.get("equity_ratio") or 0)>=40 else "Weak")),
        ("Debt / Equity",     _fmt(f.get("debt_to_equity"),"×"),"≤1.5×",   "Low" if (f.get("debt_to_equity") or 99)<=1 else ("Moderate" if (f.get("debt_to_equity") or 99)<=2 else "High")),
        ("Profitability", None, None, None),
        ("ROE",               _fmt(f.get("roe"),"%"),            "≥10%",    "Strong" if (f.get("roe") or 0)>=15 else "Weak"),
        ("ROA",               _fmt(f.get("roa"),"%"),            "≥5%",     "Strong" if (f.get("roa") or 0)>=10 else "Weak"),
        ("Operating Margin",  _fmt(f.get("operating_margin"),"%"),"≥10%",  "Strong" if (f.get("operating_margin") or 0)>=20 else "Weak"),
        ("Coverage", None, None, None),
        ("Interest Coverage", _fmt(f.get("interest_coverage"),"×"),"≥3×",  "Strong" if (f.get("interest_coverage") or 0)>=5 else "Weak"),
        ("DSCR",              _fmt(f.get("dscr"),"×"),           "≥1.5×",   "Strong" if (f.get("dscr") or 0)>=2 else "Weak"),
    ]
    rt = doc.add_table(rows=len(ratio_rows)+1, cols=4); rt.style = "Table Grid"
    for j, h in enumerate(["Metric","Value","Benchmark","Assessment"]):
        rt.rows[0].cells[j].text = h
    _header_row(rt)
    for i, (m, v, b, a) in enumerate(ratio_rows, 1):
        if v is None:
            # Category header row — set bg + bold via explicit add_run (cell.text= may leave empty run list)
            for c in rt.rows[i].cells:
                _cell_bg(c, "C7D9F5")
            para = rt.rows[i].cells[0].paragraphs[0]
            para.clear()
            para.add_run(m).bold = True
        else:
            rt.rows[i].cells[0].text = m
            rt.rows[i].cells[1].text = v or "N/A"
            rt.rows[i].cells[2].text = b or ""
            rt.rows[i].cells[3].text = a or ""
            if i % 2 == 0:
                for c in rt.rows[i].cells: _cell_bg(c, "EBF0FA")

    # ── 6. Risk Assessment ────────────────────────────────────────────────────
    doc.add_page_break()
    doc.add_heading("6. Risk Assessment", 1)
    if dec.get("strengths"):
        doc.add_heading("6.1 Strengths", 2); _bullets(doc, dec["strengths"])
    if dec.get("risks"):
        doc.add_paragraph(); doc.add_heading("6.2 Risk Factors", 2); _bullets(doc, dec["risks"])

    # ── 7. Lending Recommendation ─────────────────────────────────────────────
    doc.add_page_break()
    doc.add_heading("7. Lending Recommendation", 1)
    pd = doc.add_paragraph(); pd.add_run("Final Decision: ").bold = True
    rd = pd.add_run(rec); rd.bold = True; rd.font.size = Pt(14); rd.font.color.rgb = rc
    doc.add_paragraph()
    _kv_table(doc, [
        ("Credit Score",          f"{score:.1f}/100 ({rating})"),
        ("Loan Requested",        _fmt(analysis.get("loan_amount_requested"), f" {curr}")),
        ("Maximum Safe Loan",     _fmt(dec.get("max_safe_loan"), f" {curr}")),
        ("Loan-to-Asset Ratio",   _fmt(dec.get("loan_to_asset_ratio"), "%")),
    ])
    if dec.get("conditions"):
        doc.add_paragraph(); doc.add_heading("7.1 Approval Conditions", 2)
        _bullets(doc, dec["conditions"])
    if dec.get("monitoring_requirements"):
        doc.add_paragraph(); doc.add_heading("7.2 Monitoring Requirements", 2)
        _bullets(doc, dec["monitoring_requirements"])

    # ── 8. Appendix ───────────────────────────────────────────────────────────
    doc.add_page_break()
    doc.add_heading("8. Appendix", 1)
    doc.add_heading("8.1 Documents Reviewed", 2)
    _bullets(doc, [d["filename"] for d in docs])
    doc.add_paragraph()
    doc.add_heading("8.2 Scoring Methodology", 2)
    _kv_table(doc, [
        ("Profitability (30)", "ROE (12), Operating Margin (10), ROA (8)"),
        ("Leverage (25)",      "Equity Ratio (15), Debt/Equity (10)"),
        ("Liquidity (20)",     "Current Ratio (10), DSCR/IC (10)"),
        ("Growth (15)",        "Revenue Growth (7.5), Profit Growth (7.5)"),
        ("Qualitative (10)",   "AI-assessed business quality"),
    ], widths=(2.5, 4.0))
    doc.add_paragraph()
    pf = doc.add_paragraph()
    pf.add_run("Generated: ").bold = True
    pf.add_run(f"{today} | Credit Analysis AI System v1.0")

    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf.getvalue()